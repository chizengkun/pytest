from PyPDF2 import PdfFileReader
from docx import Document
from docx.shared import Inches
import tqdm

document = Document()
 
document.add_heading('pandas', 0)


infn = 'pandas.pdf'
pdf = PdfFileReader(open(infn,'rb'))
for i in tqdm.trange(100):   	
	document.add_paragraph( pdf.getPage(i).extractText())
document.save('demo.docx')
# def readPDF(infile):
	
	#print('documentInfo', pdf.documentInfo)
	#print('FormTextFields', pdf.getFormTextFields())
	#print('NumPages',pdf.getNumPages())	

	# p = document.add_paragraph('A plain paragraph having some ')
	# p.add_run('bold').bold = True
	# p.add_run(' and some ')
	# p.add_run('italic.').italic = True
 
	# document.add_heading('Heading, level 1', level=1)
	# document.add_paragraph('Intense quote', style='IntenseQuote')
 
	# document.add_paragraph(
 #    	'first item in unordered list', style='ListBullet'
	# )
	# document.add_paragraph(
 #    	'first item in ordered list', style='ListNumber'
	# )
 
	# document.add_picture('monty-truth.png', width=Inches(1.25))
 
	# table = document.add_table(rows=1, cols=3)
	# hdr_cells = table.rows[0].cells
	# hdr_cells[0].text = 'Qty'
	# hdr_cells[1].text = 'Id'
	# hdr_cells[2].text = 'Desc'
	# for item in recordset:
	#     row_cells = table.add_row().cells
	#     row_cells[0].text = str(item.qty)
	#     row_cells[1].text = str(item.id)
	#     row_cells[2].text = item.desc
	 
	#document.add_page_break()	 
	


